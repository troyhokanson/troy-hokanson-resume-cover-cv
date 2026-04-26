"""
Troy J. Hokanson — Locked DOCX Header Module
=============================================

SINGLE SOURCE OF TRUTH for the navy/gold header used on every DOCX
resume, cover letter, and CV. Matches the UHG reference PDF exactly:

  - Full-bleed navy #0D1B2A bar, ZERO whitespace above (page header part)
  - "Troy J. Hokanson" in WHITE Garamond-Bold, mixed case, centered
  - Thin gold #C9A84C horizontal rule, INSET (not edge-to-edge)
  - Single gold contact row beneath, pipe-separated
  - NO subtitle / role title between name and contact row
  - Section headings: steel-blue with gold underline rule

Usage in any build script:

    from docx import Document
    from templates.docx_header import (
        build_navy_header, add_section_heading, add_bullet,
        add_job_block, set_run, set_paragraph_format,
        BODY_FONT, NAME_FONT, NAVY, GOLD, STEEL, BLACK, GRAY, WHITE,
    )

    doc = Document()
    build_navy_header(doc)
    add_section_heading(doc, "Professional Summary")
    ...

DO NOT hand-roll the header in build scripts. If the locked spec needs to
change, change it HERE and only HERE.

Locked April 2026.
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ============================================================
# BRAND CONSTANTS — DO NOT CHANGE WITHOUT VERSION BUMP
# ============================================================

NAVY = RGBColor(0x0D, 0x1B, 0x2A)
GOLD = RGBColor(0xC9, 0xA8, 0x4C)
STEEL = RGBColor(0x2D, 0x6A, 0x9F)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
BLACK = RGBColor(0x14, 0x14, 0x14)
GRAY = RGBColor(0x55, 0x55, 0x55)

BODY_FONT = "Calibri"
NAME_FONT = "Garamond"  # Garamond-Bold for the name + section headings

NAME = "Troy J. Hokanson"
# Each entry: (display_text, url_or_None). url=None means plain text (no link).
CONTACT_PARTS = [
    ("Lakeville, MN", None),
    # tel: format kept simple (no +1, no formatting) to survive Word's link
    # validator on round-trip saves. Word strips tel:+1xxx links on some installs.
    ("612.352.8647", "tel:6123528647"),
    ("TroyHokanson@iCloud.com", "mailto:TroyHokanson@iCloud.com"),
    ("linkedin.com/in/troyhokanson", "https://www.linkedin.com/in/troyhokanson"),
    ("Investigative Portfolio", "https://troyhokanson.com"),
]


# ============================================================
# LOW-LEVEL XML HELPERS
# ============================================================

def shade_cell(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def shade_paragraph(paragraph, color_hex):
    """Apply background fill to a paragraph (paints behind the text and across the line height)."""
    p_pr = paragraph._p.get_or_add_pPr()
    # Remove any existing shd
    for old in p_pr.findall(qn("w:shd")):
        p_pr.remove(old)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    p_pr.append(shd)


def set_cell_margins(cell, top=100, bottom=100, left=140, right=140):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = OxmlElement("w:tcMar")
    for m, v in (("top", top), ("left", left), ("bottom", bottom), ("right", right)):
        node = OxmlElement(f"w:{m}")
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")
        tc_mar.append(node)
    tc_pr.append(tc_mar)


def remove_cell_borders(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "nil")
        tc_borders.append(b)
    tc_pr.append(tc_borders)


def add_paragraph_bottom_border(paragraph, color_hex="C9A84C", size=6, space=1):
    """Adds a bottom border (horizontal rule) to a paragraph."""
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), str(space))
    bottom.set(qn("w:color"), color_hex)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def set_run(run, *, font=BODY_FONT, size=10.5, bold=False, italic=False, color=BLACK):
    """Apply font + size + weight + color to a run, including East Asian and CS slots."""
    run.font.name = font
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    for attr in ("ascii", "hAnsi", "cs", "eastAsia"):
        rFonts.set(qn(f"w:{attr}"), font)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color


def add_hyperlink(paragraph, text, url, *, color=None, font=BODY_FONT, size=10, bold=False):
    """Append a real, clickable hyperlink to a paragraph.

    Creates a w:hyperlink element backed by a relationship in the parent part,
    so it works in Word, in PDFs converted via LibreOffice, and is preserved
    through ATS systems. Supports http(s), mailto:, and tel: URLs.
    """
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), font)
    rFonts.set(qn("w:hAnsi"), font)
    rFonts.set(qn("w:cs"), font)
    rPr.append(rFonts)

    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(int(size * 2)))
    rPr.append(sz)
    szCs = OxmlElement("w:szCs")
    szCs.set(qn("w:val"), str(int(size * 2)))
    rPr.append(szCs)

    if bold:
        b = OxmlElement("w:b")
        rPr.append(b)

    if color is not None:
        c = OxmlElement("w:color")
        c.set(qn("w:val"), "{:02X}{:02X}{:02X}".format(color[0], color[1], color[2]))
        rPr.append(c)

    new_run.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    t.set(qn("xml:space"), "preserve")
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink


def set_paragraph_format(p, *, before=0, after=0, line=1.15):
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line


# ============================================================
# THE LOCKED HEADER — the only function that should ever build it
# ============================================================

def build_navy_header(doc, *, body_top_margin_inches=1.55,
                      body_bottom_margin_inches=0.55,
                      body_left_margin_inches=0.6,
                      body_right_margin_inches=0.6):
    """
    Build the locked full-bleed navy/gold header in the page header part.

    The navy bar lives in the section.header (not the body) so it can sit
    flush at the top of the page with zero whitespace above it. The table
    is given a negative left indent equal to the page margin so it bleeds
    edge-to-edge horizontally as well.

    Layout:
        Row 1: Name in white Garamond-Bold ~28pt, centered
        Row 2: Inset gold horizontal rule (paragraph with side margins)
        Row 3: Gold contact line, centered, pipe separators
    """
    section = doc.sections[0]
    section.top_margin = Inches(body_top_margin_inches)
    section.bottom_margin = Inches(body_bottom_margin_inches)
    section.left_margin = Inches(body_left_margin_inches)
    section.right_margin = Inches(body_right_margin_inches)
    section.header_distance = Inches(0)  # bar starts at the top edge

    header = section.header
    header.is_linked_to_previous = False

    # Wipe default header paragraphs so we start clean
    for p in list(header.paragraphs):
        p_el = p._p
        p_el.getparent().remove(p_el)

    page_w = section.page_width  # full page width including margins
    page_w_twips = page_w.emu // 635  # 1 twip = 635 EMU
    margin_twips = int(section.left_margin.emu // 635)
    # Make the table WIDER than the page so it bleeds past both edges
    # (eliminates the thin white sliver on left and right when viewed in Word).
    bleed_twips = 360  # ~0.25" of extra bleed past each side
    tbl_total_twips = page_w_twips + (bleed_twips * 2)

    tbl = header.add_table(rows=1, cols=1, width=page_w)
    tbl.autofit = False
    tbl.allow_autofit = False

    tblPr = tbl._tbl.tblPr

    for old in tblPr.findall(qn("w:tblW")):
        tblPr.remove(old)
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"), str(tbl_total_twips))
    tblW.set(qn("w:type"), "dxa")
    tblPr.append(tblW)

    # Negative left indent = (-left margin - extra bleed) so the table
    # starts beyond the left page edge and extends beyond the right.
    for old in tblPr.findall(qn("w:tblInd")):
        tblPr.remove(old)
    tblInd = OxmlElement("w:tblInd")
    tblInd.set(qn("w:w"), str(-(margin_twips + bleed_twips)))
    tblInd.set(qn("w:type"), "dxa")
    tblPr.append(tblInd)

    for old in tblPr.findall(qn("w:tblLayout")):
        tblPr.remove(old)
    tblLayout = OxmlElement("w:tblLayout")
    tblLayout.set(qn("w:type"), "fixed")
    tblPr.append(tblLayout)

    cell = tbl.cell(0, 0)
    shade_cell(cell, "0D1B2A")
    remove_cell_borders(cell)
    # Tight inner padding (top a bit larger so name sits down from top edge)
    set_cell_margins(cell, top=260, bottom=240, left=200, right=200)

    tcPr = cell._tc.get_or_add_tcPr()
    tcW = tcPr.find(qn("w:tcW"))
    if tcW is None:
        tcW = OxmlElement("w:tcW")
        tcPr.append(tcW)
    tcW.set(qn("w:w"), str(tbl_total_twips))
    tcW.set(qn("w:type"), "dxa")

    # ---- Row 1: Name (white Garamond-Bold, centered, navy bg) ----
    name_p = cell.paragraphs[0]
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_format(name_p, before=0, after=0, line=1.0)
    shade_paragraph(name_p, "0D1B2A")
    r = name_p.add_run(NAME)
    set_run(r, font=NAME_FONT, size=28, bold=True, color=WHITE)

    # ---- Row 2: Inset gold horizontal rule (rendered as gold characters
    # over a NAVY-shaded paragraph so no white shows through) ----
    rule_p = cell.add_paragraph()
    rule_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_format(rule_p, before=4, after=4, line=1.0)
    shade_paragraph(rule_p, "0D1B2A")
    rule_run = rule_p.add_run("\u2500" * 78)
    set_run(rule_run, font=BODY_FONT, size=8, color=GOLD)

    # ---- Row 3: Gold contact line, centered, pipe-separated, navy bg ----
    # Each linkable piece is a real w:hyperlink so it's clickable in Word/PDF/ATS.
    contact_p = cell.add_paragraph()
    contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_format(contact_p, before=2, after=0, line=1.15)
    shade_paragraph(contact_p, "0D1B2A")
    sep = "  |  "
    for i, (text, url) in enumerate(CONTACT_PARTS):
        if i > 0:
            sep_run = contact_p.add_run(sep)
            set_run(sep_run, font=BODY_FONT, size=8.5, color=GOLD)
        if url:
            add_hyperlink(contact_p, text, url, color=GOLD, font=BODY_FONT, size=8.5)
        else:
            r = contact_p.add_run(text)
            set_run(r, font=BODY_FONT, size=8.5, color=GOLD)


# ============================================================
# SHARED BODY HELPERS — every build script should use these
# ============================================================

def add_section_heading(doc, text):
    """Steel-blue section heading with gold underline rule."""
    p = doc.add_paragraph()
    set_paragraph_format(p, before=8, after=2, line=1.1)
    run = p.add_run(text.upper())
    set_run(run, size=11.5, bold=True, color=STEEL)
    add_paragraph_bottom_border(p, color_hex="C9A84C", size=6)
    return p


def add_bullet(doc, text, *, size=10.5):
    p = doc.add_paragraph(style="List Bullet")
    set_paragraph_format(p, before=0, after=2, line=1.15)
    pf = p.paragraph_format
    pf.left_indent = Inches(0.18)
    pf.first_line_indent = Inches(-0.18)
    for r in p.runs:
        r.text = ""
    run = p.add_run(text)
    set_run(run, size=size, color=BLACK)
    return p


def add_job_block(doc, title, employer_line, dates):
    """Gold job title + italic employer line + gray dates."""
    p1 = doc.add_paragraph()
    set_paragraph_format(p1, before=6, after=0, line=1.1)
    r = p1.add_run(title)
    set_run(r, size=11, bold=True, color=GOLD)

    p2 = doc.add_paragraph()
    set_paragraph_format(p2, before=0, after=2, line=1.1)
    r = p2.add_run(employer_line)
    set_run(r, size=10.5, italic=True, color=BLACK)
    r2 = p2.add_run("    " + dates)
    set_run(r2, size=10, color=GRAY)


def new_document():
    """Returns a Document with the default Normal style set to body font."""
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = BODY_FONT
    style.font.size = Pt(10.5)
    return doc


__all__ = [
    "NAVY", "GOLD", "STEEL", "WHITE", "BLACK", "GRAY",
    "BODY_FONT", "NAME_FONT", "NAME", "CONTACT_PARTS",
    "build_navy_header",
    "add_section_heading", "add_bullet", "add_job_block",
    "set_run", "set_paragraph_format", "add_hyperlink",
    "shade_cell", "set_cell_margins", "remove_cell_borders",
    "add_paragraph_bottom_border",
    "new_document",
]
